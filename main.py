# coding: utf-8

'''
@File    :   $main.py
@Contact :   798412226@qq.com
@License :   (C)Copyright 2020-2025, HUAYUN-NLPR-CASIA

@Modify Time      @Author    @Version    @Desciption
------------      -------    --------    -----------
$2020.3.25         guozikun     1.0         None
'''
from PyQt5.QtWidgets import QMainWindow, QApplication, QDialog
from child_untitled_1 import *
from untitled import *
from datetime import timedelta
import os, time, sys, re
import matplotlib.pyplot as plt
import pandas as pd
import datetime
import numpy as np
import pymysql
import globalvar as gl
import cryptography
import chinese as ch
import matplotlib
import os
from docx import Document
from docx.shared import Inches
import importlib
importlib.reload(sys)
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号
"""
gl.set_value('globalvar_Missing', 0)
gl.set_value('globalvar_uncertainty', 0)
"""
g_Missing = 0
g_uncertainty = 0
qc = []

def config_INIT_():
    gl._init()  # 初始化全局变量管理模块
    gl.set_value('globalvar_flog', 0)# 引用全局变量管理模块 globalvar_flog 作为读入天气数据文件中的日志条数的变量

    flog = gl.get_value('globalvar_flog')

    file = open('config.cfg', mode='r+', encoding='UTF-8')
    str_config = file.read()
    gl.set_value('globalvar_config', str_config)
    gl.get_value('globalvar_config')
    file.close()



def App__RUN__():
    """
    主函数
    :return:
    """
    """设置支持高分辨率屏幕自适应"""
    QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
    """设置支持字体高分辨率自适应"""
    QApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
    app = QApplication(sys.argv)
    window = Main_windows()
    window.show()
    sys.exit(app.exec_())

class child_windows(QDialog, Ui_Dialog):
    def __init__(self):
        super(child_windows, self).__init__()
        self.setupUi(self)
        self.pushButton.clicked.connect(self.open_dir)

    def open_dir(self):
        os.system("start explorer D:\software\PyCharm Community Edition 2019.3.3\project\li_ping5.1\报告")  # c:为要打开c盘


class Main_windows(QMainWindow, Ui_MainWindow):  # 如果你是用Widget创建的窗口，这里会不同
    def __init__(self):
        super(Main_windows, self).__init__()
        self.setupUi(self)
        self.pushButton_2.clicked.connect(self.Save_datebase)
        self.pushButton_3.clicked.connect(self.DB_Search)
        self.pushButton_4.clicked.connect(self.Creat_Report)
        self.pushButton_5.clicked.connect(self.config_write)
        self.pushButton_6.clicked.connect(self.config_show)
        self.pushButton_7.clicked.connect(self.clean_win)

    def clean_win(self):
        self.textEdit_2.setText('')
    def config_write(self):
        str_config = self.textEdit_3.toPlainText()
        file = open('config.cfg', mode='r+', encoding='UTF-8')
        file.truncate()
        if str_config.find(',') == -1:
            file.close()
            self.textEdit_3.append('配置文件格式有问题 保存失败 请清空此窗口后输入配置文件')
        else:
            file.write(str_config)
            file.close()
            self.textEdit_3.append('保存成功')
    def config_show(self):
        file = open('config.cfg', mode='r+', encoding='UTF-8')
        str_config = file.read()
        self.textEdit_3.setText(str_config)
        file.close()

    def Read_dd_2(self):
        Edit_datetime = self.dateTimeEdit_2.text().replace('/', '-', 1) + ':00'
        Edit_datetime = Edit_datetime.replace(':', '+', 1)
        Edit_dict = {'Year': Edit_datetime[0:4],
                     'Mon': Edit_datetime[(Edit_datetime.index('-') + 1):(Edit_datetime.index('/'))],
                     'Day': Edit_datetime[(Edit_datetime.index('/') + 1):(Edit_datetime.index(' '))],
                     'Hour': Edit_datetime[(Edit_datetime.index(' ') + 1):(Edit_datetime.index('+'))],
                     'Min': Edit_datetime[(Edit_datetime.index('+') + 1):(Edit_datetime.index(':'))],
                     'Sec': '00'}
        if len(Edit_dict['Mon']) == 1:
            Edit_dict['Mon'] = '0' + Edit_dict['Mon']
        if len(Edit_dict['Day']) == 1:
            Edit_dict['Day'] = '0' + Edit_dict['Day']
        if len(Edit_dict['Hour']) == 1:
            Edit_dict['Hour'] = '0' + Edit_dict['Hour']
        if len(Edit_dict['Min']) == 1:
            Edit_dict['Min'] = '0' + Edit_dict['Min']
        return datetime.datetime.strptime(
            Edit_dict['Year'] + Edit_dict['Mon'] + Edit_dict['Day'] + Edit_dict['Hour'] + Edit_dict['Min'] + \
            Edit_dict['Sec'], "%Y%m%d%H%M%S")

    def Read_config(self):
        file = open('config.cfg', mode='r+', encoding='UTF-8')
        str_config = file.read()
        file.close()
        return str_config.strip().split(',')

    def Read_dd(self):
        """
        读取页面输入的时间
        :return: datetime
        """
        Edit_datetime = self.dateTimeEdit.text().replace('/', '-', 1) + ':00'
        Edit_datetime = Edit_datetime.replace(':', '+', 1)
        Edit_dict = {'Year': Edit_datetime[0:4],
                     'Mon': Edit_datetime[(Edit_datetime.index('-') + 1):(Edit_datetime.index('/'))],
                     'Day': Edit_datetime[(Edit_datetime.index('/') + 1):(Edit_datetime.index(' '))],
                     'Hour': Edit_datetime[(Edit_datetime.index(' ') + 1):(Edit_datetime.index('+'))],
                     'Min': Edit_datetime[(Edit_datetime.index('+') + 1):(Edit_datetime.index(':'))],
                     'Sec': '00'}
        if len(Edit_dict['Mon']) == 1:
            Edit_dict['Mon'] = '0' + Edit_dict['Mon']
        if len(Edit_dict['Day']) == 1:
            Edit_dict['Day'] = '0' + Edit_dict['Day']
        if len(Edit_dict['Hour']) == 1:
            Edit_dict['Hour'] = '0' + Edit_dict['Hour']
        if len(Edit_dict['Min']) == 1:
            Edit_dict['Min'] = '0' + Edit_dict['Min']
        return  datetime.datetime.strptime(
            Edit_dict['Year'] + Edit_dict['Mon'] + Edit_dict['Day'] + Edit_dict['Hour'] + Edit_dict['Min'] + \
            Edit_dict['Sec'], "%Y%m%d%H%M%S")

    def Handle_datetime(self , dateTimeEdit):
        Edit_datetime = dateTimeEdit.replace('/', '-', 1) + ':00'
        Edit_datetime = Edit_datetime.replace(':', '+', 1)
        Edit_dict = {'Year': Edit_datetime[0:4],
                     'Mon': Edit_datetime[(Edit_datetime.index('-') + 1):(Edit_datetime.index('/'))],
                     'Day': Edit_datetime[(Edit_datetime.index('/') + 1):(Edit_datetime.index(' '))],
                     'Hour': Edit_datetime[(Edit_datetime.index(' ') + 1):(Edit_datetime.index('+'))],
                     'Min': Edit_datetime[(Edit_datetime.index('+') + 1):(Edit_datetime.index(':'))],
                     'Sec': '00'}
        if len(Edit_dict['Mon']) == 1:
            Edit_dict['Mon'] = '0' + Edit_dict['Mon']
        if len(Edit_dict['Day']) == 1:
            Edit_dict['Day'] = '0' + Edit_dict['Day']
        if len(Edit_dict['Hour']) == 1:
            Edit_dict['Hour'] = '0' + Edit_dict['Hour']
        if len(Edit_dict['Min']) == 1:
            Edit_dict['Min'] = '0' + Edit_dict['Min']
        return datetime.datetime.strptime(
            Edit_dict['Year'] + Edit_dict['Mon'] + Edit_dict['Day'] + Edit_dict['Hour'] + Edit_dict['Min'] + \
            Edit_dict['Sec'], "%Y%m%d%H%M%S")


    def Str_Compare(self , str_line):
        if str(str_line[2:3]).find(self.lineEdit.text()) and str(str_line[8:9]).find(self.lineEdit_2.text()):
            return 1
        else :
            if str(str_line[2:3]).find(self.lineEdit.text()):
                self.textEdit_2.append("未匹配到ID号")
            if str(str_line[8:9]).find(self.lineEdit_2.text()):
                self.textEdit_2.append("未匹配到区站号")
            return 0

    def DI_check(self , str_line ):
        global flog
        a = ''.join(str_line[7:8])
        b = self.comboBox.currentText()
        if a == b:
            return 1
        else:
            self.textEdit_2.append("第"+str(flog)+"未找DI")
            return 0

    def ID_ckeck(self, str_line ):
        global folg
        a =  ''.join(str_line[8:9])
        b = self.lineEdit_2.text()
        if a == b:
            return 1
        else:
            self.textEdit_2.append("第"+str(flog)+"未找ID")
            return 0

    def frame_check(self, str_line ):
        global flog
        a = ''.join(str_line[10:11])
        b =  self.comboBox_2.currentText()
        if a == b:
            return 1
        else:
            self.textEdit_2.append("第"+str(flog)+"未找数据帧")
            return 0

    def StatNum_check(self, str_line ):
        a = ''.join(str_line[2:3])
        b = self.lineEdit.text()
        if a == b:
            return 1
        else:
            self.textEdit_2.append("第"+str(flog)+"未找到台站号")
            return 0

    def datetime_check(self , str_line  , dd_jure , dd_inter ):
        global dd
        global dd_pull
        global dd_Miss
        global flog
        global dd_first
        global dd_jurefirst
        dd_2 = self.Handle_dd_2()
        #if flog == 0:
            #dd_jurefirst = self.Handle_dd_jure(str_line)
        if dd <= dd_jure and dd<=dd_2 :
            if int((dd_jure - dd).seconds/60)%dd_inter == 0 :
                dd = dd_jure
                return 1
            else :
                while dd < dd_jure:
                    dd = dd +datetime.timedelta(minutes=dd_inter)
                return 0

        else:
            return 0
    def Handle_dd_jure(self , str_line):
        return datetime.datetime.strptime(''.join(str_line[9:10]), '%Y%m%d%H%M%S')
    def read_asNULL(self , str_line):
        global flog
        g_MessageDict[str(flog)] = ['NULL']
        flog = flog+1
        g_MessageDict[str(flog)] = str_line[9:10]
    def readfile_asline(self , str_line ):
        global g_MessageLine
        g_MessageLine = g_MessageLine+str_line[2:3]+str_line[7:8]+ str_line[8:9]+ str_line[10:11]+ str_line[9:10]+str_line[10:]
        return

    def Read_combox_3(self):
        a = int(self.comboBox_3.currentIndex())
        if a == 0:
            return 1
        if a == 1:
            return 5
        if a == 2:
            return 10
        if a == 3:
            return 60

    def Handle_MessageLine(self):
        global g_MessageLine
        MessageLine = g_MessageLine
        dd = self.Handle_dd()
        dd_2 = self.Handle_dd_2()
        dd_len = int((dd_2 - dd).seconds/60)
        i = 0

        while i <= dd_len:
            if datetime.datetime.strptime(str(MessageLine[i]), '%Y%m%d%H%M%S') == dd:
                dd = dd + datetime.timedelta(minutes=1)
                i = i+1
            else:
                dd_Message = datetime.datetime.strptime(str(MessageLine[i]), '%Y%m%d%H%M%S')
                list_len = int((dd_Message - dd).seconds/60)
                list = [0]*list_len
                dd = datetime.datetime.strptime(MessageLine[i], '%Y%m%d%H%M%S')
                MessageLine[i:i] = list
                i = i + list_len

        return MessageLine

    def Creat_Table(self, str_line, mycursor, conn, Table_Name):
        """
        :param str_line:
        :param mycursor:
        :param conn:
        :param Table_Name:
        :return:
        """
        check = mycursor.execute("show table status like %s", Table_Name)
        conn.commit()
        if check == 0:
            sql = """CREATE TABLE `wetherdate`.`%s` (
            `id` INT NOT NULL AUTO_INCREMENT,
            `datetime` DATETIME(1) NOT NULL,
            `date` TEXT(1000) NOT NULL,
            PRIMARY KEY (`id`),
            UNIQUE INDEX `datetime_UNIQUE` (`datetime` ASC) VISIBLE);"""%(Table_Name)
            mycursor.execute(sql)

        return

    def Table_to_sql(self, sql, Table_Name):
        """
        :param sql:
        :param Table_Name:
        :return: void
        """
        sql = sql.replace('TABLE_NAME', Table_Name, 1)
        return sql

    def save_SQL_asline(self, line , str_line, mycursor, conn):
        global flog
        Table_Name = str(str_line[2] + '_' + str_line[7] + '_' + str_line[8] + '_' + str_line[10])
        time = datetime.datetime.strptime(str(str_line[9]), '%Y%m%d%H%M%S')
        self.Creat_Table(str_line, mycursor, conn, Table_Name)
        sql = """INSERT INTO `wetherdate`.`TABLE_NAME` ( datetime, date) VALUES (%s, %s)"""
        sql = self.Table_to_sql(sql, Table_Name)
        val = (time, line)
        # 执行sql语句
        try:
            mycursor.execute(sql, val)
            #提交到数据库
            conn.commit()
            flog = flog +1
            print(flog)
        except:
            flog = flog + 1
            self.textEdit_2.append('ERROR' + str(flog))
            print('ERROR'+str(flog))

    def Save_datebase(self):
        """
        This is the trigger function of the button (stored in the database) on the interface
        return : void
        parameters : No parameters required
        Author : guozikun
        """
        #dd_inter = datetime.datetime.strptime(str(self.Read_combox_3()), "%S")
        dd_inter = self.Read_combox_3()
        global dd
        global dd_2
        global g_MessageDict
        global flog
        global dd_jurefirst
        global dd_first
        global End_identification
        End_identification = 0
        file_name = self.lineEdit_3.text()
        dd = self.Handle_datetime(self.dateTimeEdit.text())
        dd_2 = self.Handle_datetime(self.dateTimeEdit_2.text())
        if len(self.lineEdit.text()) == 0 or len(self.lineEdit_2.text()) == 0:
            if len(self.lineEdit.text()) == 0:
                self.textEdit_2.append("提示：请输入区站号")
            if len(self.lineEdit_2.text()) == 0:
                self.textEdit_2.append("提示：请输入ID号")

        else:
            if os.path.isfile(file_name):
                self.textEdit_2.append("验证文件成功")
                # 打开数据库连接-填入你Mysql的账号密码和端口
                conn = pymysql.connect('localhost', 'root', '2667885', "wetherdate", charset='utf8')
                # 使用 cursor() 方法创建一个游标对象 cursor
                mycursor = conn.cursor()
                file = open(file_name, mode='r+', encoding='UTF-8')
                flog = 0
                for line in file.readlines() :
                    if len(line) != 0:
                        str_line = line.strip().split(',')
                        if len(str_line) >= 11:
                            if self.Str_Compare(str_line):
                                dd_jure = self.Handle_dd_jure(str_line)
                                if ( self.ID_ckeck(str_line )) and (self.DI_check(str_line )) and (self.frame_check(str_line)) and (self.StatNum_check(str_line )):
                                    self.save_SQL_asline(line, str_line, mycursor, conn)

                            else:
                                break
                file.close()
                mycursor.close()
                conn.close()

            else:
                self.textEdit.append('未找到文件，请放到根目录下')

    def abnormal_exist(self, qc_1, qc_miss, qc_8, qc_2):
        error = 1
        if len(qc_1) != 0:
            return error
        if len(qc_miss) != 0:
            return error
        if len(qc_8) != 0:
            return error
        if len(qc_2) != 0:
            return error
        return 0
    def Printinfo_picture(self, checkbox_position , picture_data, qc_data, num_data, num_dataloss):
        """

        :param checkbox_position:
        :param picture_data:
        :param qc_data:
        :param num_data:
        :param num_dataloss:
        :return:
        """
        """重置画布"""
        plt.clf()

        error = 2
        miss = 8
        uncertain = 1
        config = self.Read_config()
        plt.title('Atmospheric data')
        plt.xlabel('Retrieve of '+str(num_data)+ ' data, data loss = '+str(num_dataloss))
        plt.ylabel('data')
        # 以0.2为间隔均匀采样
        len_X = int((((self.Read_dd_2()) - (self.Read_dd())).seconds/60) + (((self.Read_dd_2()) - (self.Read_dd())).days*1440))

        for i in range(len(picture_data)):
            """画曲线"""
            list_data = [float(j) for j in picture_data[i]]
            plt.plot(list(list_data), '.', markersize=1.5, label= str((config[checkbox_position[i]])))
            """画数据丢失的点"""
            plt.plot(self.get_Missing_position(picture_data[i], qc_data[i]),
                     [0] * len(self.get_Missing_position(picture_data[i], qc_data[i])),
                     'o', label='数据丢失 ' + str(num_dataloss))
            """画qc = 8 缺测 """
            plt.plot(self.get_measuring_position(picture_data[i], qc_data[i], miss),
                     [0] * len(self.get_measuring_position(picture_data[i], qc_data[i], miss)), 'o',
                     label='缺测 ' + str(
                         len(self.get_measuring_position(picture_data[i], qc_data[i], miss))))
            """qc = 1 存疑"""
            plt.plot(self.get_position_x(picture_data[i], qc_data[i], uncertain),
                     self.get_position_y(picture_data[i], qc_data[i], uncertain), 'o',
                     label='存疑 '+str(len(self.get_measuring_position(picture_data[i], qc_data[i], uncertain))))
            """画qc == 2 错误"""
            plt.plot(self.get_position_x(picture_data[i], qc_data[i], error),
                     self.get_position_y(picture_data[i], qc_data[i], error), 'o',
                     label='错误 '+str(len(self.get_measuring_position(picture_data[i], qc_data[i], error))))

            """
            输出窗口提示信息
            """
            """#这个判断作用是是否有异常点，有异常点救出提示，无异常点就不输出了
            if self.abnormal_exist(self.get_measuring_position(picture_data[i], qc_data[i], miss),
                                   self.get_Missing_position(picture_data[i], qc_data[i]),
                                   self.get_position_x(picture_data[i], qc_data[i], uncertain),
                                   self.get_position_x(picture_data[i], qc_data[i], error)) == 1:"""

            self.textEdit_2.append(str(config[checkbox_position[i]]) + '质控统计:')
            if len(self.get_measuring_position(picture_data[i], qc_data[i], miss)) != 0:
                self.textEdit_2.append(
                    '    缺测  ' +str(len(self.get_measuring_position(picture_data[i], qc_data[i], miss))) )
            else :
                self.textEdit_2.append('    缺测  0')
            if len(self.get_position_x(picture_data[i], qc_data[i], uncertain)) != 0:
                self.textEdit_2.append('    存疑  ' +str(len(self.get_position_x(picture_data[i], qc_data[i], uncertain))))
            else:
                self.textEdit_2.append('    存疑  0')
            if len(self.get_position_x(picture_data[i], qc_data[i], error)) != 0:
                self.textEdit_2.append('    错误  ' +str(len(self.get_position_x(picture_data[i], qc_data[i], error))) )
            else:
                self.textEdit_2.append('    错误  0')

        plt.legend(bbox_to_anchor=(1.01, 1), loc=2, borderaxespad=0., handleheight=1.675)
        plt.savefig('testblueline.jpg', dpi=200, bbox_inches='tight')
        plt.show()

    def Real_ele(self, ele):
        if ele == 'AAA':
            return '温度'
        if ele == 'AB10':
            return '地温1层'
        if ele == 'AB20':
            return '地温2层'
        if ele == 'AB30':
            return '地温3层'
        if ele == 'AB40':
            return '地温4层'
        if ele == 'AB50':
            return '地温5层'
        if ele == 'ADA':
            return '湿度'
        if ele == 'AB10':
            return '地温1层'
        if ele == 'AFA':
            return '10米风'
        if ele == 'AFA150':
            return '1.5米风'
        if ele == 'AGA':
            return '气压'
        if ele == 'AHA':
            return '翻斗雨'
        if ele == 'AHC':
            return '称重雨量'
        if ele == 'AJA':
            return '辐射'
        if ele == 'ARG10':
            return '土壤水分1层'
        if ele == 'ARG20':
            return '土壤水分2层'
        if ele == 'ARG30':
            return '土壤水分3层'
        if ele == 'ARG40':
            return '土壤水分4层'
        if ele == 'ARG50':
            return '土壤水分5层'
        return '未找到对应元素'




    def Save_picture(self, doc, checkbox_position , picture_data, qc_data, picture_name, num_data, num_dataloss):
        """

        :param doc:
        :param checkbox_position:
        :param picture_data:
        :param qc_data:
        :param picture_name:
        :param num_data:
        :param num_dataloss:
        :return:
        """
        """重置画布"""
        plt.clf()

        error = 2
        miss = 8
        uncertain = 1
        config = self.Read_config()
        plt.title('Atmospheric data')
        plt.xlabel('Retrieve of ' + str(num_data) + ' data, data loss = ' + str(num_dataloss))
        plt.ylabel('data')
        # 以0.2为间隔均匀采样
        len_X = int((((self.Read_dd_2()) - (self.Read_dd())).seconds / 60) + (
                    ((self.Read_dd_2()) - (self.Read_dd())).days * 1440))

        for i in range(len(picture_data)):
            """如果是气压，数据都除10000"""
            if config[checkbox_position[i]] == 'AGA':
                a = [float(q)/10000 for q in picture_data[i]]
                picture_data[i] = a
            """画曲线"""
            list_data = [float(j) for j in picture_data[i]]
            plt.plot(list(list_data), '-', label=str((config[checkbox_position[i]])))
            """画数据丢失的点"""
            plt.plot(self.get_Missing_position(picture_data[i], qc_data[i]),
                     [0] * len(self.get_Missing_position(picture_data[i], qc_data[i])),
                     'o', label='数据丢失 ' + str(num_dataloss))
            """画qc = 8 缺测 """
            plt.plot(self.get_measuring_position(picture_data[i], qc_data[i], miss),
                     [0] * len(self.get_measuring_position(picture_data[i], qc_data[i], miss)), 'o',
                     label='缺测 ' + str(
                         len(self.get_measuring_position(picture_data[i], qc_data[i], miss))))
            """qc = 1 存疑"""
            plt.plot(self.get_position_x(picture_data[i], qc_data[i], uncertain),
                     self.get_position_y(picture_data[i], qc_data[i], uncertain), 'o',
                     label='存疑 ' + str(len(self.get_measuring_position(picture_data[i], qc_data[i], uncertain))))
            """画qc == 2 错误"""
            plt.plot(self.get_position_x(picture_data[i], qc_data[i], error),
                     self.get_position_y(picture_data[i], qc_data[i], error), 'o',
                     label='错误 ' + str(len(self.get_measuring_position(picture_data[i], qc_data[i], error))))

            """
            输出窗口提示信息
            """
            """#这个判断作用是是否有异常点，有异常点救出提示，无异常点就不输出了
            if self.abnormal_exist(self.get_measuring_position(picture_data[i], qc_data[i], miss),
                                   self.get_Missing_position(picture_data[i], qc_data[i]),
                                   self.get_position_x(picture_data[i], qc_data[i], uncertain),
                                   self.get_position_x(picture_data[i], qc_data[i], error)) == 1:"""

            doc.add_paragraph(self.Real_ele(str(config[checkbox_position[i]])) + '统计:')
            doc.add_paragraph('    数据丢失  ' + str(num_dataloss))
            if len(self.get_measuring_position(picture_data[i], qc_data[i], miss)) != 0:
                doc.add_paragraph(
                    '    缺测  ' + str(len(self.get_measuring_position(picture_data[i], qc_data[i], miss))))
            else:
                doc.add_paragraph('    缺测  0')
            if len(self.get_position_x(picture_data[i], qc_data[i], uncertain)) != 0:
                doc.add_paragraph(
                    '    存疑  ' + str(len(self.get_position_x(picture_data[i], qc_data[i], uncertain))))
            else:
                doc.add_paragraph('    存疑  0')
            if len(self.get_position_x(picture_data[i], qc_data[i], error)) != 0:
                doc.add_paragraph(
                    '    错误  ' + str(len(self.get_position_x(picture_data[i], qc_data[i], error))))
            else:
                doc.add_paragraph('    错误  0')

        plt.legend(bbox_to_anchor=(1.01, 1), loc=2, borderaxespad=0., handleheight=1.675)
        plt.savefig(picture_name, dpi=200, bbox_inches='tight')


    def Chackbox(self):
        checkbox_state = []

        for i in range(1,len(self.Read_config())):
            self.temp = getattr(self, "checkBox_%d" % i)
            if self.temp.isChecked():
                checkbox_state.append(1)
            else:
                checkbox_state.append(0)
        return checkbox_state


    """def Creat_Struct_date(self):
        file = open('config.cfg', mode='r+', encoding='UTF-8')
        str_config = file.read().strip().split(',')
        file.close()
        Date_type = np.dtype([('row_date', 'S300'),
                              ('row_qc', 'S300'),
                              ('QC_8', 'I'),
                              ('QC_2', 'I'),
                              ('QC_MISS', 'I')])
        Struct_date = np.array([('THIS IS BRGIN',
                                 'THIS IS BEGIN',
                                 0,
                                 0,
                                 0,)] * len(str_config),
                               dtype=Date_type)
        return Struct_date"""


    def Read_specif_ele(self, results, loop_1, checkbox_position):
        """
        读取所有的元素数值大小返回列表
        :param results:
        :param loop_1:
        :return: a = ['0154', '0155', '0156', '0155']
        """
        global  g_Missing
        global g_uncertainty
        global qc
        g_Missing = 0
        g_uncertainty = 0
        a = []
        config = self.Read_config()
        for row in results:
            str_line = str(row[2]).strip().split(',')[13:]
            a.append(str_line[checkbox_position[loop_1]*2+1])
            qc.append((str_line[(len(config))*2][checkbox_position[loop_1]]))
            if str_line[(len(config))*2][checkbox_position[loop_1]] == '1':
                g_uncertainty = g_uncertainty + 1
            if str_line[(len(config))*2][checkbox_position[loop_1]] == '8':
                g_Missing = g_Missing + 1

        """self.textEdit_2.append('   缺测：'+str(g_Missing)+'    存疑:' + str(g_uncertainty))"""

        return a

    def get_Checkstatus_position(self, checkbox_state):
        """
        用来存放checkbox的选择位置
        :param checkbox_state:
        :return: checkbox_state[2, 6, 45, 78, .............]
        """
        a = []
        checkbox_position = list(enumerate(checkbox_state))
        for i in checkbox_position:
            if i [1] == 1:
                a.append(i[0])
        return a

    def get_Missing_position(self, state, qc):
        """
        用来存放checkbox的选择位置
        :param checkbox_state:
        :return: checkbox_state[2, 6, 45, 78, .............]
        """
        a = []
        checkbox_position = list(enumerate(state))
        for i in range(len(checkbox_position)):
            if checkbox_position[i][1] == '0' and qc[i] == 'N':
                a.append(checkbox_position[i][0])
        return a

    def get_measuring_position(self, state, qc, miss):
        """
        用来存放checkbox的选择位置
        :param checkbox_state:
        :return: checkbox_state[2, 6, 45, 78, .............]
        """
        a = []
        checkbox_position = list(enumerate(state))
        for i in range(len(checkbox_position)):
            if qc[i] == str(miss):
                a.append(checkbox_position[i][0])
        return a


    def get_position_x(self, state, qc, qc_para):
        """
        用来存放checkbox的选择位置
        :param checkbox_state:
        :return: checkbox_state[2, 6, 45, 78, .............]
        """
        a = []
        checkbox_position = list(enumerate(state))
        for i in range(len(checkbox_position)):
            if  qc[i] == str(qc_para):
                a.append(checkbox_position[i][0])
        return a

    def get_position_y(self, state, qc, qc_para):
        """
        用来存放checkbox的选择位置
        :param checkbox_state:
        :return: checkbox_state[2, 6, 45, 78, .............]
        """
        a = []
        checkbox_position = list(enumerate(state))
        for i in range(len(checkbox_position)):
            if  qc[i] == str(qc_para):
                a.append(int(checkbox_position[i][1]))
        return a


    def Read_inter(self):
        a = self.comboBox_3.currentIndex()
        if a == 0:
            return 1
        if a == 1:
            return 5
        if a == 2:
            return 10
        if a == 3:
            return 60

    def Read_specif_qc(self, results, loop_1, checkbox_position):
        qc = []
        config = self.Read_config()
        for row in results:
            str_line = str(row[2]).strip().split(',')[13:]
            qc.append((str_line[(len(config))*2][checkbox_position[loop_1]]))
        return qc

    def Repair_result(self, result):
        i = 0
        results = []
        #delta = timedelta(minutes = 1)
        dd = self.Read_dd()
        if len(result) != 0:
            dd = result [i][1]
        dd_2 = self.Read_dd_2()

        while dd <= dd_2:
            if i < len(result):
                if result [i][1] == dd:
                    results.append(result [i])
                    i = i + 1
                else:
                    results.append(('id',dd, 'BG,001,57495,394827,1162815,00444,14,YIIP,0,datetime,001,043,03,AAA,\
0,AAA5i,0,AB10,0,AB20,0,AB30,0,AB40,0,AB50,0,ADA,0,ADB,0,AEA,0,AEA150,0,AEB,0,\
AEB150,0,AEC,0,AEC150,0,AED,0,AED150,0,AEF,0,AEF150,0,AFA,0,AFA150,0,AFA150a,0,AFAa,\
0,AFB,0,AFB150,0,AFC,0,AFC150,0,AFD,0,AFD150,0,AGA,0,AHA,0,AHA5,00,AHC,0,AHC5,00,\
AJA,0,AJAa,0,AJAc,0,AJT,201911051005,ARG10,0,ARG20,0,ARG30,0,ARG40,0,ARG50,0,\
NNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNN,z,1,rL,1,xA,7,9748,ED'))
            else:
                results.append(('id', dd, 'BG,001,57495,394827,1162815,00444,14,YIIP,0,datetime,001,043,03,AAA,\
0,AAA5i,0,AB10,0,AB20,0,AB30,0,AB40,0,AB50,0,ADA,0,ADB,0,AEA,0,AEA150,0,AEB,0,\
AEB150,0,AEC,0,AEC150,0,AED,0,AED150,0,AEF,0,AEF150,0,AFA,0,AFA150,0,AFA150a,0,AFAa,\
0,AFB,0,AFB150,0,AFC,0,AFC150,0,AFD,0,AFD150,0,AGA,0,AHA,0,AHA5,00,AHC,0,AHC5,00,\
AJA,0,AJAa,0,AJAc,0,AJT,201911051005,ARG10,0,ARG20,0,ARG30,0,ARG40,0,ARG50,0,\
NNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNN,z,1,rL,1,xA,7,9748,ED'))

            dd = dd + datetime.timedelta(minutes = self.Read_inter())

        """if result[len(result)-1][1] != dd_2:
            results.append(('id',dd,'id',dd, 'BG,001,57495,394827,1162815,00444,14,YIIP,0,datetime,001,043,03,AAA,\
0,AAA5i,0,AB10,0,AB20,0,AB30,0,AB40,0,AB50,0,ADA,0,ADB,0,AEA,0,AEA150,0,AEB,0,\
AEB150,0,AEC,0,AEC150,0,AED,0,AED150,0,AEF,0,AEF150,0,AFA,0,AFA150,0,AFA150a,0,AFAa,\
0,AFB,0,AFB150,0,AFC,0,AFC150,0,AFD,0,AFD150,0,AGA,0,AHA,0,AHA5,00,AHC,0,AHC5,00,\
AJA,0,AJAa,0,AJAc,0,AJT,201911051005,ARG10,0,ARG20,0,ARG30,0,ARG40,0,ARG50,0,\
NNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNNN,z,1,rL,1,xA,7,9748,ED'))"""

        return results

        """
        while i < len(results) - 1 and len(results) > 1:
            Time_apart = (results[i+1][1] - results[i][1])
            if (results[i+1][1] - results[i][1]) != timedelta(minutes = 1):
                for loop in range(int((results[i+1][1] - results[i][1]).seconds/60)-1):
                    results[i+1:i+1] = [0,0,0]
                    i = i+1"""

        return tuple(results)

    def printinfo_MissingNum(self, results):
        a = int((self.Read_dd_2() - self.Read_dd()).seconds / 60) + 1
        b = len(results)
        self.textEdit_2.append('数据缺失：'  + str(a - b)+ '条')

    def Replace_result(self, results_):
        """

        :param results:
        :return: results
        """
        replaced_result = []
        b = ()
        for i in range(len(results_)):
            a = list(results_[i])

            if a[2].find('/') != -1:
                for j in range(len(str(a[2]))):
                    if a[2][j].find('/') != -1:
                        a[2] = str(a[2]).replace('/', '0')


                replaced_result.append(tuple(a))
            else:
                replaced_result.append(tuple(a))
                pass

        return replaced_result


    def DB_Search(self):
        global  g_Missing
        global g_uncertainty
        picture_date = []
        picture_qc = []
        check_num = 0
        Table_Name = self.lineEdit.text() + '_' + self.comboBox.currentText()+ '_' + self.lineEdit_2.text() + '_' + self.comboBox_2.currentText()
        beg_time = self.Read_dd()
        end_time = self.Read_dd_2()
        mydb = pymysql.connect(
            host="localhost",
            user="root",
            passwd="2667885",
            database="wetherdate"
        )
        mycursor = mydb.cursor()
        sql = "SELECT * FROM TABLE_NAME WHERE datetime >= '%s' and datetime <='%s'"%(beg_time, end_time)
        sql = self.Table_to_sql(sql, Table_Name)
        mycursor.execute(sql)
        # fetchall() 获取所有记录
        #Struct_date = self.Creat_Struct_date()
        db_data = mycursor.fetchall()
        repare_data = self.Repair_result(db_data)
        results = self.Replace_result(repare_data)

        #窗口提示信息
        self.textEdit_2.append(
            '++++++++++++++++共检索' + str(len(results)) + '条数据,其中数据缺失'+ str(((int((self.Read_dd_2() - self.Read_dd()).days * 1440) + int((self.Read_dd_2() - self.Read_dd()).seconds / 60)+1 - len(db_data)))) + '条' +'++++++++++++++++++++++++')

        """获得checkbox页面的勾选的原始状态"""
        checkbox_state = self.Chackbox()
        """提取checkbox勾选的位置到列表"""
        checkbox_position = self.get_Checkstatus_position(checkbox_state)

        for loop in range(len(checkbox_state)):
            if checkbox_state[loop] == 1:
                check_num = check_num + 1
        picture_date = []
        picture_qc = []

        for loop_1 in range(check_num):

            data = self.Read_specif_ele(results, loop_1, checkbox_position)
            exec('list_'+str(loop_1)+'='+str(data))
            qc_data = self.Read_specif_qc(results, loop_1, checkbox_position)
            exec('qc_' + str(loop_1) + '=' + str(qc_data))
            #print('list_' + str(loop_1) + ':', eval('list_' + str(loop_1)))
            """self.textEdit_2.append(str('list_' + str(loop_1) + ':') + str(eval('list_' + str(loop_1))))"""
            """self.textEdit_2.append(str('qc_' + str(loop_1) + ':') + str(eval('qc_' + str(loop_1))))"""

            picture_date.append(tuple(eval('list_' + str(loop_1))))

            picture_qc.append(tuple(eval('qc_' + str(loop_1))))

        mycursor.close()
        mydb.close()
        self.Printinfo_picture(checkbox_position ,
                               picture_date,
                               picture_qc,
                               num_data = str(len(results)),
                               num_dataloss = ((int((self.Read_dd_2() - self.Read_dd()).days * 1440) + int((self.Read_dd_2() - self.Read_dd()).seconds / 60)+1 - len(db_data)))  )

        """self.child = child_windows()#
        self.child = wingdows()
        self.child.show()"""

        """
        Missing_num = 
        x = range(100)
        y = np.sin(x)
        t = np.cos(x)
        plt.plot(x, y, ls="-", lw=2, label="plot figure")
        plt.plot(x, t, label="t")
        plt.show()
        
        """
    def Name_datetime(self, str_datetime):
        str_datetime = str_datetime.replace(' ', '_')
        str_datetime = str_datetime.replace(':', '_')
        return str_datetime


    def Creat_Report(self):
        """
        此函数为按钮生成报告的函数
        :return:
        """
        global g_Missing
        global g_uncertainty
        doc = Document()
        picture_date = []
        picture_qc = []
        check_num = 0
        Table_Name = self.lineEdit.text() + '_' + self.comboBox.currentText() + '_' + self.lineEdit_2.text() + '_' + self.comboBox_2.currentText()
        beg_time = self.Read_dd()
        end_time = self.Read_dd_2()
        mydb = pymysql.connect(
            host="localhost",
            user="root",
            passwd="2667885",
            database="wetherdate"
        )
        mycursor = mydb.cursor()
        sql = "SELECT * FROM TABLE_NAME WHERE datetime >= '%s' and datetime <='%s'" % (beg_time, end_time)
        sql = self.Table_to_sql(sql, Table_Name)
        mycursor.execute(sql)
        # fetchall() 获取所有记录
        # Struct_date = self.Creat_Struct_date()
        db_data = mycursor.fetchall()
        repare_data = self.Repair_result(db_data)
        results = self.Replace_result(repare_data)

        """生成需要写入报告的提示信息"""
        str_word = (
            '                共检索' + str(len(results)) + '条数据,其中数据缺失' + str(((
                        int((self.Read_dd_2() - self.Read_dd()).days * 1440) + int(
                    (self.Read_dd_2() - self.Read_dd()).seconds / 60) + 1 - len(
                    db_data)))) + '条' + '                    ')
        """添加文字到docx"""
        doc.add_paragraph(str_word)
        doc.add_paragraph('时间：' + str((self.Read_dd())) + '  致  ' + str((self.Read_dd_2())))

        """生成checkbox勾选的位置到列表"""
        checkbox_position = []

        checkbox_position.append([0])
        checkbox_position.append([2,3,4,5,6])
        checkbox_position.append([7])
        checkbox_position.append([9,10])
        checkbox_position.append([29])
        checkbox_position.append([30])
        checkbox_position.append([32])
        checkbox_position.append([34])
        checkbox_position.append([38,39,40,41,42])


        for i in range(len(checkbox_position)):
            """以选择了的checkbox位置信息来命名图片"""
            picture_name = str(checkbox_position[i]) + '.jpg'
            picture_date = []
            picture_qc = []
            for loop_1 in range(len(checkbox_position[i])):
                data = self.Read_specif_ele(results, loop_1, checkbox_position[i])
                exec('list_' + str(loop_1) + '=' + str(data))
                qc_data = self.Read_specif_qc(results, loop_1, checkbox_position[i])
                exec('qc_' + str(loop_1) + '=' + str(qc_data))
                # print('list_' + str(loop_1) + ':', eval('list_' + str(loop_1)))
                """self.textEdit_2.append(str('list_' + str(loop_1) + ':') + str(eval('list_' + str(loop_1))))"""
                """self.textEdit_2.append(str('qc_' + str(loop_1) + ':') + str(eval('qc_' + str(loop_1))))"""

                picture_date.append(tuple(eval('list_' + str(loop_1))))

                picture_qc.append(tuple(eval('qc_' + str(loop_1))))


            """把图像保存成jpg文件"""
            self.Save_picture(doc,
                             checkbox_position[i],
                             picture_date,
                             picture_qc,
                             picture_name,
                             num_data = str(len(results)),
                             num_dataloss = ((int((self.Read_dd_2() - self.Read_dd()).days * 1440) + int((
                                            self.Read_dd_2() - self.Read_dd()).seconds / 60) + 1 - len(db_data)))
                             )
            """把图片存入doc"""
            doc.add_picture(picture_name, width=Inches(6))
        """关闭数据库"""
        mycursor.close()
        mydb.close()

        """添加图, 设置宽度"""
        doc.save('报告\\' + self.lineEdit.text() +'站' + str(self.Name_datetime(str(self.Read_dd())))+'至' + str(self.Name_datetime(str(self.Read_dd_2())))+'报告'+'.docx')

        self.child = child_windows()
        self.child.show()



if __name__ == "__main__":

    config_INIT_()
    App__RUN__()
